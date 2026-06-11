"""
resume_parser.py
----------------
Extracts structured information from resume PDF/DOCX/TXT files.
Works for ANY domain — tech, MBA, finance, healthcare, etc.
"""

import pdfplumber
import spacy
import re
import os
import docx

nlp = spacy.load("en_core_web_sm")

# ── Name blacklist ─────────────────────────────────────────────────────────────
BLACKLIST_NAMES = [
    "dart", "java", "python", "swift", "go", "rust",
    "ruby", "c", "c++", "c#", "sql", "html", "css", "r"
]

# ── Universal skills list ──────────────────────────────────────────────────────
UNIVERSAL_SKILLS = [
    # Programming Languages
    "python", "java", "javascript", "typescript", "c++", "c#", "c", "r",
    "scala", "go", "rust", "swift", "kotlin", "dart", "php", "ruby",
    "matlab", "bash", "shell", "perl", "vba",
    # Web & Mobile
    "react", "angular", "vue", "nodejs", "node.js", "django", "flask",
    "fastapi", "spring", "flutter", "android", "ios", "html", "css",
    "next.js", "express", "tailwind", "bootstrap", "graphql", "rest api",
    "redux", "jquery", "webpack", "sass",
    # Data & ML
    "machine learning", "deep learning", "nlp", "computer vision",
    "data science", "data analysis", "data engineering", "statistics",
    "pandas", "numpy", "scikit-learn", "tensorflow", "keras", "pytorch",
    "matplotlib", "seaborn", "tableau", "power bi", "excel", "sql",
    "hadoop", "spark", "etl", "a/b testing", "regression",
    "classification", "neural network", "opencv", "mediapipe",
    "hugging face", "transformers", "langchain",
    # Cloud & DevOps
    "aws", "azure", "gcp", "docker", "kubernetes", "ci/cd", "git",
    "github", "gitlab", "linux", "terraform", "jenkins", "ansible",
    "microservices", "agile", "scrum", "jira",
    # Databases
    "mongodb", "mysql", "postgresql", "firebase", "redis", "cassandra",
    "dynamodb", "oracle", "sqlite", "elasticsearch", "supabase",
    # Business & Management
    "project management", "product management", "business analysis",
    "stakeholder management", "kanban", "strategy", "consulting",
    "operations", "supply chain", "logistics", "business development",
    "market research", "competitive analysis",
    # Marketing & Sales
    "digital marketing", "seo", "sem", "social media", "content marketing",
    "email marketing", "google analytics", "crm", "salesforce", "hubspot",
    "brand management", "marketing strategy", "sales", "lead generation",
    "b2b", "b2c", "performance marketing",
    # Finance & Accounting
    "financial modeling", "financial analysis", "accounting", "budgeting",
    "forecasting", "valuation", "investment banking", "equity research",
    "portfolio management", "risk management", "audit", "taxation",
    "ifrs", "gaap", "tally", "sap", "bloomberg", "cfa", "ca",
    "mergers and acquisitions", "due diligence",
    # HR
    "recruitment", "talent acquisition", "onboarding", "payroll",
    "performance management", "employee relations", "hr analytics",
    "learning and development", "compensation", "hris", "workday",
    # Design
    "figma", "adobe xd", "sketch", "photoshop", "illustrator",
    "ui design", "ux design", "user research", "wireframing",
    "prototyping", "design thinking", "canva", "after effects",
    # Healthcare & Biotech
    "clinical research", "pharmacology", "biotechnology", "microbiology",
    "biochemistry", "genomics", "pcr", "elisa", "cell culture",
    "regulatory affairs", "clinical trials", "bioinformatics",
    # Cybersecurity
    "penetration testing", "vulnerability assessment", "nmap", "wireshark",
    "burp suite", "metasploit", "siem", "wazuh", "splunk", "owasp",
    "ethical hacking", "network security", "cryptography", "firewall",
    # Tools
    "microsoft office", "powerpoint", "word", "outlook", "notion",
    "slack", "trello", "asana", "confluence", "postman", "vs code",
    "jupyter", "android studio", "xcode",
    # Soft Skills
    "leadership", "communication", "teamwork", "problem solving",
    "critical thinking", "time management", "negotiation", "presentation",
    "analytical skills", "decision making",
]

# ── Skill stopwords ────────────────────────────────────────────────────────────
SKILL_STOPWORDS = {
    "hands", "core", "inc", "brands", "market", "award", "growth",
    "tools", "learning", "university", "tool", "insight", "analyst",
    "passion", "solution", "interface", "responsibilities", "location",
    "delivery", "part", "creation", "business", "cloud", "data",
    "team", "work", "years", "experience", "skills", "role", "global",
    "strong", "good", "great", "high", "new", "key", "well", "ability",
    "knowledge", "candidate", "looking", "seeking", "required",
    "preferred", "must", "minimum", "including", "etc", "like",
    "using", "use", "based", "related", "across", "within", "between",
    "please", "apply", "position", "company", "opportunity", "join"
}


# ── Extract text from file ─────────────────────────────────────────────────────
def extract_text(file_path):
    """Extract text from PDF, DOCX, or TXT files."""
    ext = os.path.splitext(file_path)[1].lower()

    if ext == ".pdf":
        full_text = ""
        try:
            with pdfplumber.open(file_path) as pdf:
                for page in pdf.pages:
                    page_text = page.extract_text()
                    if page_text:
                        full_text += page_text + "\n"
        except Exception as e:
            print(f"  PDF read error: {e}")
        return full_text.strip()

    elif ext == ".docx":
        full_text = ""
        try:
            doc = docx.Document(file_path)
            for para in doc.paragraphs:
                if para.text.strip():
                    full_text += para.text + "\n"
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        if cell.text.strip():
                            full_text += cell.text + "\n"
        except Exception as e:
            print(f"  DOCX read error: {e}")
        return full_text.strip()

    elif ext == ".txt":
        try:
            with open(file_path, "r", encoding="utf-8", errors="ignore") as f:
                return f.read().strip()
        except Exception as e:
            print(f"  TXT read error: {e}")
            return ""

    else:
        print(f"  Unsupported file type: {ext}")
        return ""


# ── Clean text ─────────────────────────────────────────────────────────────────
def clean_text(text):
    """Fix squished words, remove junk, normalize whitespace."""
    text = re.sub(r'([a-z])([A-Z])',      r'\1 \2', text)
    text = re.sub(r'([A-Z])([A-Z][a-z])', r'\1 \2', text)
    text = re.sub(r',([a-z])',             r', \1',  text)
    text = re.sub(r'\band([A-Za-z])',      r'and \1', text)
    text = re.sub(r'\bthe([A-Z])',         r'the \1', text)
    text = re.sub(r'\bof([A-Z])',          r'of \1',  text)
    text = re.sub(r'\(cid:\d+\)',          '',        text)
    text = re.sub(r'\s+',                  ' ',       text)
    text = re.sub(r'[^\w\s@.\-+,/]',      ' ',       text)
    return text.strip()


# ── Extract skills ─────────────────────────────────────────────────────────────
def extract_skills(text):
    """
    Match text against universal skills list.
    Uses word boundary matching to avoid false positives.
    e.g. won't match 'scala' inside 'scalable'.
    """
    text_lower = text.lower()
    found      = []

    for skill in UNIVERSAL_SKILLS:
        if skill in SKILL_STOPWORDS:
            continue

        if ' ' in skill:
            # Multi-word skill — exact substring match
            if skill in text_lower:
                found.append(skill)
        else:
            # Single word — use word boundary to avoid partial matches
            pattern = r'(?<![a-z])' + re.escape(skill) + r'(?![a-z])'
            if re.search(pattern, text_lower):
                found.append(skill)

    return found


# ── Match JD skills against resume ────────────────────────────────────────────
def match_skills(resume_text, jd_skills):
    """
    Match JD skills against resume text.
    Uses exact word boundary matching.
    Returns found and missing skills.
    """
    resume_lower   = resume_text.lower()
    real_jd_skills = [
        s for s in jd_skills
        if len(s) > 2 and s not in SKILL_STOPWORDS
    ]

    found   = []
    missing = []

    for skill in real_jd_skills:
        if ' ' in skill:
            matched = skill in resume_lower
        else:
            pattern = r'(?<![a-z])' + re.escape(skill) + r'(?![a-z])'
            matched = bool(re.search(pattern, resume_lower))

        if matched:
            found.append(skill)
        else:
            missing.append(skill)

    return found, missing


# ── Extract contact info ───────────────────────────────────────────────────────
def extract_contact_info(raw_text, cleaned_text):
    """Extract name, email, and phone from resume."""

    # Use first non-empty line as name — most reliable for resumes
    lines = [l.strip() for l in raw_text.split('\n') if l.strip()]
    name  = ""

    if lines:
        first_line = lines[0].strip()
        # Remove phone numbers and emails from first line if any
        first_line = re.sub(r'[\+\d][\d\s\-]{7,}', '', first_line)
        first_line = re.sub(r'[\w.\-]+@[\w.\-]+\.\w+', '', first_line)
        first_line = first_line.strip()
        if len(first_line) > 2:
            name = first_line

    # Fallback to spaCy if first line looks wrong
    if not name or len(name.split()) > 5 or any(c.isdigit() for c in name):
        doc = nlp(raw_text[:500])
        for ent in doc.ents:
            if ent.label_ == "PERSON":
                candidate = ent.text.strip()
                if candidate.lower() not in BLACKLIST_NAMES:
                    name = candidate
                    break

    email = re.findall(r'[\w.\-]+@[\w.\-]+\.\w+', cleaned_text)
    phone = re.findall(r'[\+\d][\d\s\-]{8,15}\d',  cleaned_text)

    return {
        "name":  name,
        "email": email[0] if email else "",
        "phone": phone[0].strip() if phone else ""
    }


# ── Extract sections ───────────────────────────────────────────────────────────
def extract_sections(raw_text):
    """Detect and extract Education, Experience, Projects sections."""
    sections = {"experience": "", "education": "", "projects": ""}
    current  = None

    for line in raw_text.split('\n'):
        l = line.lower().strip()

        if any(w in l for w in ["experience", "work history",
                                  "employment", "internship"]):
            current = "experience"
        elif any(w in l for w in ["education", "qualification", "academic"]):
            current = "education"
        elif any(w in l for w in ["project", "projects"]):
            current = "projects"
        elif current:
            line = re.sub(r'([a-z])([A-Z])', r'\1 \2', line)
            line = re.sub(r'\(cid:\d+\)',     '',       line)
            sections[current] += line.strip() + " "

    for key in sections:
        sections[key] = sections[key].strip()[:500]

    return sections


# ── Main parse function ────────────────────────────────────────────────────────
def parse_resume(file_path):
    """
    Main function — takes file path (PDF/DOCX/TXT),
    returns structured dictionary.
    """
    if not os.path.exists(file_path):
        return {"error": f"File not found: {file_path}"}

    raw = extract_text(file_path)

    if not raw:
        return {"error": "No text extracted — file may be corrupted or image-based"}

    cleaned  = clean_text(raw)
    contact  = extract_contact_info(raw, cleaned)
    skills   = extract_skills(cleaned)
    sections = extract_sections(raw)

    return {
        "name":       contact["name"],
        "email":      contact["email"],
        "phone":      contact["phone"],
        "skills":     skills,
        "experience": sections["experience"],
        "education":  sections["education"],
        "projects":   sections["projects"],
        "full_text":  cleaned
    }